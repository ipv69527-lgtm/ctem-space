from __future__ import annotations

import html
import re
from datetime import datetime
from pathlib import Path

from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.models.asset import Asset
from app.models.report import Report
from app.models.sync_task import SyncTask
from app.models.template import Template
from app.models.unit import Unit
from app.models.vulnerability import Vulnerability

REPORT_TYPE_ASSET = "资产清单报告"
REPORT_TYPE_VULN = "漏洞整改报告"
REPORT_TYPE_RISK = "单位风险排行报告"
REPORT_TYPE_SYNC = "同步质量报告"

LEGACY_REPORT_TYPE_MAP = {
    "资产漏洞清单": REPORT_TYPE_ASSET,
    "单位整改报告": REPORT_TYPE_VULN,
}

REMEDIATION_STATUSES = {"待整改", "整改中", "待复测"}
RISK_WEIGHT = {"严重": 10, "高危": 7, "中危": 4, "低危": 1}


def _report_type(report: Report) -> str:
    return LEGACY_REPORT_TYPE_MAP.get(report.type, report.type)


async def generate_report_file(
    db: AsyncSession,
    report: Report,
    template: Template | None = None,
    severity_filter: list[str] | None = None,
    status_filter: list[str] | None = None,
) -> str:
    report_dir = Path(settings.REPORT_DIR)
    report_dir.mkdir(parents=True, exist_ok=True)
    severity_filter = severity_filter or []
    status_filter = status_filter or []

    unit = None
    if report.unit_id:
        unit_result = await db.execute(select(Unit).where(Unit.id == report.unit_id))
        unit = unit_result.scalar_one_or_none()
    units_stmt = select(Unit)
    if report.unit_id:
        units_stmt = units_stmt.where(Unit.id == report.unit_id)
    units = list((await db.execute(units_stmt)).scalars().all())

    asset_stmt = select(Asset)
    if report.unit_id:
        asset_stmt = asset_stmt.where(Asset.unit_id == report.unit_id)
    assets = list((await db.execute(asset_stmt)).scalars().all())
    asset_ids = {asset.id for asset in assets}

    vulns = list((await db.execute(select(Vulnerability))).scalars().all())
    if report.unit_id:
        vulns = [v for v in vulns if asset_ids.intersection(set(v.asset_ids or []))]
    if severity_filter:
        vulns = [v for v in vulns if v.severity in severity_filter]
    if status_filter:
        vulns = [v for v in vulns if v.status in status_filter]

    unit_ids = {u.id for u in units}
    task_stmt = select(SyncTask)
    if unit_ids:
        task_stmt = task_stmt.where(SyncTask.unit_id.in_(unit_ids))
    sync_tasks = list((await db.execute(task_stmt.order_by(SyncTask.created_at.desc()))).scalars().all())

    file_name = f"{report.id}.{report.format}"
    file_path = report_dir / file_name

    if report.format == "docx":
        _write_docx(file_path, report, unit, units, assets, vulns, sync_tasks, template)
    elif report.format == "xlsx":
        _write_xlsx(file_path, report, unit, units, assets, vulns, sync_tasks, template)
    elif report.format == "html":
        _write_html(file_path, report, unit, units, assets, vulns, sync_tasks, template)
    else:
        raise ValueError("不支持的报表格式")

    return str(file_path)


def _fmt_dt(value) -> str:
    return value.isoformat(sep=" ", timespec="seconds") if value else "-"


def _count_by(items: list, attr: str) -> dict[str, int]:
    counts: dict[str, int] = {}
    for item in items:
        key = str(getattr(item, attr, "") or "-")
        counts[key] = counts.get(key, 0) + 1
    return counts


def _unit_stats(units: list[Unit], assets: list[Asset], vulns: list[Vulnerability], sync_tasks: list[SyncTask]) -> list[dict]:
    assets_by_unit: dict[str, list[Asset]] = {}
    for asset in assets:
        assets_by_unit.setdefault(asset.unit_id, []).append(asset)

    tasks_by_unit: dict[str, list[SyncTask]] = {}
    for task in sync_tasks:
        tasks_by_unit.setdefault(task.unit_id, []).append(task)

    rows = []
    for unit in units:
        unit_assets = assets_by_unit.get(unit.id, [])
        asset_ids = {asset.id for asset in unit_assets}
        unit_vulns = [v for v in vulns if asset_ids.intersection(set(v.asset_ids or []))]
        critical_high = sum(1 for v in unit_vulns if v.severity in {"严重", "高危"})
        remediation = sum(1 for v in unit_vulns if v.status in REMEDIATION_STATUSES)
        score = len(unit_assets) + sum(RISK_WEIGHT.get(v.severity, 1) for v in unit_vulns)
        latest_task = tasks_by_unit.get(unit.id, [None])[0]
        rows.append({
            "unit": unit,
            "asset_count": len(unit_assets),
            "vuln_count": len(unit_vulns),
            "critical_high": critical_high,
            "remediation": remediation,
            "score": score,
            "last_sync": unit.last_sync,
            "last_task_status": latest_task.status if latest_task else "",
            "last_task_message": latest_task.message if latest_task else "",
        })
    rows.sort(key=lambda item: item["score"], reverse=True)
    return rows


def _sync_quality_rows(units: list[Unit], sync_tasks: list[SyncTask]) -> list[dict]:
    tasks_by_unit: dict[str, list[SyncTask]] = {}
    for task in sync_tasks:
        tasks_by_unit.setdefault(task.unit_id, []).append(task)
    rows = []
    for unit in units:
        tasks = tasks_by_unit.get(unit.id, [])
        success_count = sum(1 for task in tasks if task.status == "success")
        failed_count = sum(1 for task in tasks if task.status == "failed")
        running_count = sum(1 for task in tasks if task.status in {"pending", "running"})
        latest_task = tasks[0] if tasks else None
        rows.append({
            "unit": unit,
            "task_count": len(tasks),
            "success_count": success_count,
            "failed_count": failed_count,
            "running_count": running_count,
            "success_rate": round(success_count / len(tasks) * 100, 1) if tasks else 0,
            "latest_task": latest_task,
        })
    rows.sort(key=lambda item: (item["failed_count"], -item["success_rate"]), reverse=True)
    return rows


def _template_context(
    report: Report,
    unit: Unit | None,
    units: list[Unit],
    assets: list[Asset],
    vulns: list[Vulnerability],
    sync_tasks: list[SyncTask],
    template: Template | None,
) -> dict[str, object]:
    report_type = _report_type(report)
    remediation_count = sum(1 for v in vulns if v.status in REMEDIATION_STATUSES)
    critical_high = sum(1 for v in vulns if v.severity in {"严重", "高危"})
    sync_failed_count = sum(1 for task in sync_tasks if task.status == "failed")
    risk_rows = _unit_stats(units, assets, vulns, sync_tasks)
    top_risk_unit = risk_rows[0]["unit"].name if risk_rows else "-"
    return {
        "report_title": report.title,
        "report_type": report_type,
        "template_name": template.name if template else "",
        "unit_name": unit.name if unit else "全量单位",
        "generated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC"),
        "asset_count": len(assets),
        "vuln_count": len(vulns),
        "critical_high": critical_high,
        "remediation_count": remediation_count,
        "unit_count": len(units),
        "sync_task_count": len(sync_tasks),
        "sync_failed_count": sync_failed_count,
        "top_risk_unit": top_risk_unit,
    }


def _render_template_content(
    report: Report,
    unit: Unit | None,
    units: list[Unit],
    assets: list[Asset],
    vulns: list[Vulnerability],
    sync_tasks: list[SyncTask],
    template: Template | None,
) -> str:
    if not template or not (template.content or "").strip():
        return ""
    context = _template_context(report, unit, units, assets, vulns, sync_tasks, template)

    def replace(match: re.Match[str]) -> str:
        key = match.group(1).strip()
        return str(context.get(key, match.group(0)))

    return re.sub(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}", replace, template.content).strip()


def _replace_placeholders(text: str, context: dict[str, object]) -> str:
    def replace(match: re.Match[str]) -> str:
        key = match.group(1).strip()
        return str(context.get(key, match.group(0)))

    return re.sub(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}", replace, text)


def _set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def _iter_docx_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                yield paragraph
            for table in container.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            yield paragraph


def _fill_docx_table(table, headers: list[str], rows: list[list[object]]) -> None:
    for idx, title in enumerate(headers):
        table.rows[0].cells[idx].text = title
    for values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = _fmt_dt(value) if isinstance(value, datetime) else str(value if value is not None and value != "" else "-")


def _append_docx_table(doc, title: str, headers: list[str], rows: list[list[object]]):
    if title:
        doc.add_heading(title, level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    _fill_docx_table(table, headers, rows)
    return table


def _insert_docx_table_after(doc, paragraph, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    _fill_docx_table(table, headers, rows)
    paragraph._p.addnext(table._tbl)


def _docx_table_payloads(
    report_type: str,
    units: list[Unit],
    assets: list[Asset],
    vulns: list[Vulnerability],
    sync_tasks: list[SyncTask],
) -> dict[str, tuple[str, list[str], list[list[object]]]]:
    asset_map = {asset.id: asset for asset in assets}
    unit_by_id = {u.id: u.name for u in units}
    return {
        "asset_table": (
            "资产清单",
            ["名称", "IP", "单位", "类型", "风险", "端口", "服务", "最近发现"],
            [
                [asset.name, asset.ip, unit_by_id.get(asset.unit_id, "-"), asset.type, asset.risk, asset.ports, asset.services, asset.last_seen]
                for asset in assets
            ],
        ),
        "vuln_table": (
            "漏洞清单",
            ["标题", "CVE", "CVSS", "等级", "状态", "影响资产", "最后发现", "描述"],
            [
                [vuln.title, vuln.cve, f"{vuln.cvss:.1f}", vuln.severity, vuln.status, _affected_assets_text(vuln, asset_map), vuln.last_found, vuln.desc]
                for vuln in vulns
            ],
        ),
        "tracking_table": (
            "处置跟踪",
            ["漏洞", "状态", "处置备注", "状态更新时间", "修复建议"],
            [[vuln.title, vuln.status, vuln.status_note, vuln.status_updated_at, vuln.solution] for vuln in vulns],
        ),
        "risk_rank_table": (
            "单位风险排行",
            ["单位", "资产数", "漏洞数", "严重/高危", "待整改", "风险评分", "最近同步", "最近任务"],
            [
                [
                    item["unit"].name,
                    item["asset_count"],
                    item["vuln_count"],
                    item["critical_high"],
                    item["remediation"],
                    item["score"],
                    item["last_sync"],
                    item["last_task_status"],
                ]
                for item in _unit_stats(units, assets, vulns, sync_tasks)
            ],
        ),
        "sync_quality_table": (
            "同步质量",
            ["单位", "任务数", "成功", "失败", "执行中", "成功率", "最近任务", "最近结果"],
            [
                [
                    item["unit"].name,
                    item["task_count"],
                    item["success_count"],
                    item["failed_count"],
                    item["running_count"],
                    f"{item['success_rate']}%",
                    item["latest_task"].status if item["latest_task"] else "-",
                    (item["latest_task"].error_detail or item["latest_task"].message) if item["latest_task"] else "-",
                ]
                for item in _sync_quality_rows(units, sync_tasks)
            ],
        ),
    }


def _append_docx_standard_tables(
    doc,
    report_type: str,
    units: list[Unit],
    assets: list[Asset],
    vulns: list[Vulnerability],
    sync_tasks: list[SyncTask],
) -> None:
    payloads = _docx_table_payloads(report_type, units, assets, vulns, sync_tasks)
    if report_type in {REPORT_TYPE_RISK, REPORT_TYPE_SYNC}:
        title, headers, rows = payloads["risk_rank_table"]
        _append_docx_table(doc, title, headers, rows)
    if report_type == REPORT_TYPE_SYNC:
        title, headers, rows = payloads["sync_quality_table"]
        _append_docx_table(doc, title, headers, rows)
    if report_type in {REPORT_TYPE_ASSET, REPORT_TYPE_VULN, REPORT_TYPE_RISK}:
        title, headers, rows = payloads["asset_table"]
        _append_docx_table(doc, title, headers, rows)
    if report_type in {REPORT_TYPE_VULN, REPORT_TYPE_RISK}:
        title, headers, rows = payloads["vuln_table"]
        _append_docx_table(doc, title, headers, rows)
        title, headers, rows = payloads["tracking_table"]
        _append_docx_table(doc, title, headers, rows)
    if report_type not in {REPORT_TYPE_ASSET, REPORT_TYPE_VULN, REPORT_TYPE_RISK, REPORT_TYPE_SYNC}:
        for key in ("asset_table", "vuln_table"):
            title, headers, rows = payloads[key]
            _append_docx_table(doc, title, headers, rows)


def _write_docx_from_template(
    path: Path,
    report: Report,
    unit: Unit | None,
    units: list[Unit],
    assets: list[Asset],
    vulns: list[Vulnerability],
    sync_tasks: list[SyncTask],
    template: Template,
) -> None:
    from docx import Document

    template_path = Path(template.file_path or "")
    if not template_path.exists() or not template_path.is_file():
        raise ValueError("Word 模板文件不存在")

    report_type = _report_type(report)
    doc = Document(template_path)
    context = _template_context(report, unit, units, assets, vulns, sync_tasks, template)
    payloads = _docx_table_payloads(report_type, units, assets, vulns, sync_tasks)
    inserted_tables: set[str] = set()

    for paragraph in list(_iter_docx_paragraphs(doc)):
        original = paragraph.text
        if not original:
            continue
        for key, (_, headers, rows) in payloads.items():
            token = "{{" + key + "}}"
            if token in original:
                original = original.replace(token, "").strip()
                _set_paragraph_text(paragraph, original)
                _insert_docx_table_after(doc, paragraph, headers, rows)
                inserted_tables.add(key)
        rendered = _replace_placeholders(original, context)
        if rendered != paragraph.text:
            _set_paragraph_text(paragraph, rendered)

    if not inserted_tables:
        _append_docx_standard_tables(doc, report_type, units, assets, vulns, sync_tasks)

    doc.save(path)


def _write_docx(
    path: Path,
    report: Report,
    unit: Unit | None,
    units: list[Unit],
    assets: list[Asset],
    vulns: list[Vulnerability],
    sync_tasks: list[SyncTask],
    template: Template | None,
) -> None:
    from docx import Document

    if template and template.file_path:
        _write_docx_from_template(path, report, unit, units, assets, vulns, sync_tasks, template)
        return

    asset_map = {asset.id: asset for asset in assets}
    report_type = _report_type(report)
    doc = Document()
    doc.add_heading(report.title, 0)
    doc.add_paragraph(f"生成时间：{datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC")
    doc.add_paragraph(f"报表范围：{unit.name if unit else '全量单位'}")
    doc.add_paragraph(f"报表类型：{report_type}")
    doc.add_heading("核心指标", level=1)
    doc.add_paragraph(f"资产数量：{len(assets)}")
    doc.add_paragraph(f"漏洞数量：{len(vulns)}")
    doc.add_paragraph(f"严重/高危漏洞：{sum(1 for v in vulns if v.severity in ['严重', '高危'])}")
    doc.add_paragraph(f"待整改/整改中/待复测：{sum(1 for v in vulns if v.status in ['待整改', '整改中', '待复测'])}")
    rendered_content = _render_template_content(report, unit, units, assets, vulns, sync_tasks, template)
    if rendered_content:
        doc.add_heading("模板正文", level=1)
        for paragraph in rendered_content.splitlines():
            doc.add_paragraph(paragraph)

    if report_type in {REPORT_TYPE_RISK, REPORT_TYPE_SYNC}:
        doc.add_heading("单位风险排行", level=1)
        table = doc.add_table(rows=1, cols=8)
        for idx, title in enumerate(("单位", "资产数", "漏洞数", "严重/高危", "待整改", "风险评分", "最近同步", "最近任务")):
            table.rows[0].cells[idx].text = title
        for item in _unit_stats(units, assets, vulns, sync_tasks):
            row = table.add_row().cells
            row[0].text = item["unit"].name
            row[1].text = str(item["asset_count"])
            row[2].text = str(item["vuln_count"])
            row[3].text = str(item["critical_high"])
            row[4].text = str(item["remediation"])
            row[5].text = str(item["score"])
            row[6].text = _fmt_dt(item["last_sync"])
            row[7].text = item["last_task_status"] or "-"

    if report_type == REPORT_TYPE_SYNC:
        doc.add_heading("同步质量", level=1)
        table = doc.add_table(rows=1, cols=8)
        for idx, title in enumerate(("单位", "任务数", "成功", "失败", "执行中", "成功率", "最近任务", "最近结果")):
            table.rows[0].cells[idx].text = title
        for item in _sync_quality_rows(units, sync_tasks):
            latest = item["latest_task"]
            row = table.add_row().cells
            row[0].text = item["unit"].name
            row[1].text = str(item["task_count"])
            row[2].text = str(item["success_count"])
            row[3].text = str(item["failed_count"])
            row[4].text = str(item["running_count"])
            row[5].text = f"{item['success_rate']}%"
            row[6].text = latest.status if latest else "-"
            row[7].text = latest.error_detail or latest.message if latest else "-"

    if report_type in {REPORT_TYPE_ASSET, REPORT_TYPE_VULN, REPORT_TYPE_RISK}:
        doc.add_heading("资产清单", level=1)
        table = doc.add_table(rows=1, cols=8)
        for idx, title in enumerate(("名称", "IP", "单位", "类型", "风险", "端口", "服务", "最近发现")):
            table.rows[0].cells[idx].text = title
        unit_by_id = {u.id: u.name for u in units}
        for asset in assets:
            row = table.add_row().cells
            row[0].text = asset.name
            row[1].text = asset.ip
            row[2].text = unit_by_id.get(asset.unit_id, "-")
            row[3].text = asset.type
            row[4].text = asset.risk
            row[5].text = asset.ports or "-"
            row[6].text = asset.services or "-"
            row[7].text = _fmt_dt(asset.last_seen)

    if report_type in {REPORT_TYPE_VULN, REPORT_TYPE_RISK}:
        doc.add_heading("漏洞清单", level=1)
        table = doc.add_table(rows=1, cols=8)
        for idx, title in enumerate(("标题", "CVE", "CVSS", "等级", "状态", "影响资产", "最后发现", "描述")):
            table.rows[0].cells[idx].text = title
        for vuln in vulns:
            row = table.add_row().cells
            row[0].text = vuln.title
            row[1].text = vuln.cve or "-"
            row[2].text = f"{vuln.cvss:.1f}"
            row[3].text = vuln.severity
            row[4].text = vuln.status
            row[5].text = _affected_assets_text(vuln, asset_map)
            row[6].text = _fmt_dt(vuln.last_found)
            row[7].text = vuln.desc or "-"

        doc.add_heading("处置跟踪", level=1)
        table = doc.add_table(rows=1, cols=5)
        for idx, title in enumerate(("漏洞", "状态", "处置备注", "状态更新时间", "修复建议")):
            table.rows[0].cells[idx].text = title
        for vuln in vulns:
            row = table.add_row().cells
            row[0].text = vuln.title
            row[1].text = vuln.status
            row[2].text = vuln.status_note or "-"
            row[3].text = _fmt_dt(vuln.status_updated_at)
            row[4].text = vuln.solution or "-"

    if report_type not in {REPORT_TYPE_ASSET, REPORT_TYPE_VULN, REPORT_TYPE_RISK, REPORT_TYPE_SYNC}:
        doc.add_heading("资产清单", level=1)
        table = doc.add_table(rows=1, cols=7)
        for idx, title in enumerate(("名称", "IP", "类型", "风险", "端口", "服务", "最近发现")):
            table.rows[0].cells[idx].text = title
        for asset in assets:
            row = table.add_row().cells
            row[0].text = asset.name
            row[1].text = asset.ip
            row[2].text = asset.type
            row[3].text = asset.risk
            row[4].text = asset.ports or "-"
            row[5].text = asset.services or "-"
            row[6].text = _fmt_dt(asset.last_seen)

        doc.add_heading("漏洞清单", level=1)
        table = doc.add_table(rows=1, cols=7)
        for idx, title in enumerate(("标题", "CVE", "CVSS", "等级", "状态", "影响资产", "最后发现")):
            table.rows[0].cells[idx].text = title
        for vuln in vulns:
            row = table.add_row().cells
            row[0].text = vuln.title
            row[1].text = vuln.cve or "-"
            row[2].text = f"{vuln.cvss:.1f}"
            row[3].text = vuln.severity
            row[4].text = vuln.status
            row[5].text = _affected_assets_text(vuln, asset_map)
            row[6].text = _fmt_dt(vuln.last_found)
    doc.save(path)

    return

    table = doc.add_table(rows=1, cols=7)
    for idx, title in enumerate(("名称", "IP", "类型", "风险", "端口", "服务", "最近发现")):
        table.rows[0].cells[idx].text = title
    for asset in assets:
        row = table.add_row().cells
        row[0].text = asset.name
        row[1].text = asset.ip
        row[2].text = asset.type
        row[3].text = asset.risk
        row[4].text = asset.ports or "-"
        row[5].text = asset.services or "-"
        row[6].text = asset.last_seen.isoformat() if asset.last_seen else "-"

    doc.add_heading("漏洞清单", level=1)
    table = doc.add_table(rows=1, cols=7)
    for idx, title in enumerate(("标题", "CVE", "CVSS", "等级", "状态", "影响资产", "最后发现")):
        table.rows[0].cells[idx].text = title
    for vuln in vulns:
        row = table.add_row().cells
        row[0].text = vuln.title
        row[1].text = vuln.cve or "-"
        row[2].text = f"{vuln.cvss:.1f}"
        row[3].text = vuln.severity
        row[4].text = vuln.status
        row[5].text = _affected_assets_text(vuln, asset_map)
        row[6].text = vuln.last_found.isoformat() if vuln.last_found else "-"

    doc.add_heading("处置跟踪", level=1)
    table = doc.add_table(rows=1, cols=5)
    for idx, title in enumerate(("漏洞", "状态", "处置备注", "状态更新时间", "修复建议")):
        table.rows[0].cells[idx].text = title
    for vuln in vulns:
        row = table.add_row().cells
        row[0].text = vuln.title
        row[1].text = vuln.status
        row[2].text = vuln.status_note or "-"
        row[3].text = vuln.status_updated_at.isoformat() if vuln.status_updated_at else "-"
        row[4].text = vuln.solution or "-"
    doc.save(path)


def _write_xlsx(
    path: Path,
    report: Report,
    unit: Unit | None,
    units: list[Unit],
    assets: list[Asset],
    vulns: list[Vulnerability],
    sync_tasks: list[SyncTask],
    template: Template | None,
) -> None:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter

    asset_map = {asset.id: asset for asset in assets}
    unit_by_id = {u.id: u.name for u in units}
    report_type = _report_type(report)
    wb = Workbook()
    ws = wb.active
    ws.title = "概览"
    rows = [
        ("报表标题", report.title),
        ("生成时间", datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")),
        ("报表范围", unit.name if unit else "全量单位"),
        ("报表类型", report_type),
        ("模板名称", template.name if template else ""),
        ("资产数量", len(assets)),
        ("漏洞数量", len(vulns)),
        ("严重/高危漏洞", sum(1 for v in vulns if v.severity in ["严重", "高危"])),
        ("待整改/整改中/待复测", sum(1 for v in vulns if v.status in ["待整改", "整改中", "待复测"])),
    ]
    for row in rows:
        ws.append(row)
    rendered_content = _render_template_content(report, unit, units, assets, vulns, sync_tasks, template)
    if rendered_content:
        ws.append(("模板正文", rendered_content))

    if report_type in {REPORT_TYPE_ASSET, REPORT_TYPE_VULN, REPORT_TYPE_RISK}:
        asset_ws = wb.create_sheet("资产清单")
        asset_ws.append(["名称", "IP", "单位", "类型", "系统", "风险", "开放端口", "服务", "位置", "运营商", "最近发现"])
        for asset in assets:
            asset_ws.append([
                asset.name,
                asset.ip,
                unit_by_id.get(asset.unit_id, ""),
                asset.type,
                asset.os,
                asset.risk,
                asset.ports,
                asset.services,
                asset.location,
                asset.isp,
                asset.last_seen,
            ])

    if report_type in {REPORT_TYPE_VULN, REPORT_TYPE_RISK}:
        vuln_ws = wb.create_sheet("漏洞清单")
        vuln_ws.append(["标题", "CVE", "CVSS", "等级", "状态", "影响资产", "影响资产数", "首次发现", "最后发现", "描述", "修复方案"])
        for vuln in vulns:
            affected_assets = _affected_assets_text(vuln, asset_map)
            vuln_ws.append([
                vuln.title,
                vuln.cve,
                vuln.cvss,
                vuln.severity,
                vuln.status,
                affected_assets,
                len([asset_id for asset_id in (vuln.asset_ids or []) if asset_id in asset_map]),
                vuln.first_found,
                vuln.last_found,
                vuln.desc,
                vuln.solution,
            ])

        track_ws = wb.create_sheet("处置跟踪")
        track_ws.append(["漏洞", "CVE", "等级", "状态", "处置备注", "状态更新时间", "影响资产", "修复建议"])
        for vuln in vulns:
            track_ws.append([
                vuln.title,
                vuln.cve,
                vuln.severity,
                vuln.status,
                vuln.status_note,
                vuln.status_updated_at,
                _affected_assets_text(vuln, asset_map),
                vuln.solution,
            ])

    if report_type in {REPORT_TYPE_RISK, REPORT_TYPE_SYNC}:
        risk_ws = wb.create_sheet("单位风险排行")
        risk_ws.append(["单位", "编码", "资产数", "漏洞数", "严重/高危", "待整改", "风险评分", "最近同步", "最近任务状态", "最近任务结果"])
        for item in _unit_stats(units, assets, vulns, sync_tasks):
            risk_ws.append([
                item["unit"].name,
                item["unit"].code,
                item["asset_count"],
                item["vuln_count"],
                item["critical_high"],
                item["remediation"],
                item["score"],
                item["last_sync"],
                item["last_task_status"],
                item["last_task_message"],
            ])

    if report_type == REPORT_TYPE_SYNC:
        sync_ws = wb.create_sheet("同步质量")
        sync_ws.append(["单位", "任务数", "成功", "失败", "执行中", "成功率", "最近状态", "最近结果", "最近更新时间", "最近查询条件"])
        for item in _sync_quality_rows(units, sync_tasks):
            latest = item["latest_task"]
            sync_ws.append([
                item["unit"].name,
                item["task_count"],
                item["success_count"],
                item["failed_count"],
                item["running_count"],
                f"{item['success_rate']}%",
                latest.status if latest else "",
                latest.error_detail or latest.message if latest else "",
                latest.updated_at if latest else None,
                latest.query_condition if latest else "",
            ])

        detail_ws = wb.create_sheet("同步任务明细")
        detail_ws.append(["单位", "任务ID", "状态", "查询条件", "拉取资产", "入库资产", "入库漏洞", "结果", "错误详情", "创建时间", "更新时间"])
        for task in sync_tasks:
            detail_ws.append([
                unit_by_id.get(task.unit_id, ""),
                task.id,
                task.status,
                task.query_condition,
                task.fetched_assets,
                task.synced_assets,
                task.synced_vulns,
                task.message,
                task.error_detail,
                task.created_at,
                task.updated_at,
            ])

    if report_type not in {REPORT_TYPE_ASSET, REPORT_TYPE_VULN, REPORT_TYPE_RISK, REPORT_TYPE_SYNC}:
        asset_ws = wb.create_sheet("资产清单")
        asset_ws.append(["名称", "IP", "类型", "系统", "风险", "开放端口", "服务", "位置", "运营商", "最近发现"])
        for asset in assets:
            asset_ws.append([
                asset.name,
                asset.ip,
                asset.type,
                asset.os,
                asset.risk,
                asset.ports,
                asset.services,
                asset.location,
                asset.isp,
                asset.last_seen,
            ])

    for sheet in wb.worksheets:
        sheet.freeze_panes = "A2"
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="595959")
        for column_cells in sheet.columns:
            length = max(len(str(cell.value or "")) for cell in column_cells)
            sheet.column_dimensions[get_column_letter(column_cells[0].column)].width = min(max(length + 2, 12), 48)
    wb.save(path)
    return

    asset_ws = wb.create_sheet("资产清单")
    asset_ws.append(["名称", "IP", "类型", "系统", "风险", "开放端口", "服务", "位置", "运营商", "最近发现"])
    for asset in assets:
        asset_ws.append([
            asset.name,
            asset.ip,
            asset.type,
            asset.os,
            asset.risk,
            asset.ports,
            asset.services,
            asset.location,
            asset.isp,
            asset.last_seen,
        ])

    vuln_ws = wb.create_sheet("漏洞清单")
    vuln_ws.append(["标题", "CVE", "CVSS", "等级", "状态", "影响资产", "影响资产数", "最后发现", "描述", "修复方案"])
    for vuln in vulns:
        affected_assets = _affected_assets_text(vuln, asset_map)
        vuln_ws.append([
            vuln.title,
            vuln.cve,
            vuln.cvss,
            vuln.severity,
            vuln.status,
            affected_assets,
            len([asset_id for asset_id in (vuln.asset_ids or []) if asset_id in asset_map]),
            vuln.last_found,
            vuln.desc,
            vuln.solution,
        ])

    track_ws = wb.create_sheet("处置跟踪")
    track_ws.append(["漏洞", "CVE", "等级", "状态", "处置备注", "状态更新时间", "影响资产", "修复建议"])
    for vuln in vulns:
        track_ws.append([
            vuln.title,
            vuln.cve,
            vuln.severity,
            vuln.status,
            vuln.status_note,
            vuln.status_updated_at,
            _affected_assets_text(vuln, asset_map),
            vuln.solution,
        ])

    for sheet in wb.worksheets:
        sheet.freeze_panes = "A2"
        for cell in sheet[1]:
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor="595959")
        for column_cells in sheet.columns:
            length = max(len(str(cell.value or "")) for cell in column_cells)
            sheet.column_dimensions[get_column_letter(column_cells[0].column)].width = min(max(length + 2, 12), 48)
    wb.save(path)


def _write_html(
    path: Path,
    report: Report,
    unit: Unit | None,
    units: list[Unit],
    assets: list[Asset],
    vulns: list[Vulnerability],
    sync_tasks: list[SyncTask],
    template: Template | None,
) -> None:
    def e(value: object) -> str:
        return html.escape("" if value is None else str(value))

    asset_map = {asset.id: asset for asset in assets}
    unit_by_id = {u.id: u.name for u in units}
    report_type = _report_type(report)
    asset_rows = "\n".join(
        f"<tr><td>{e(a.name)}</td><td>{e(a.ip)}</td><td>{e(unit_by_id.get(a.unit_id, '-'))}</td><td>{e(a.type)}</td><td>{e(a.risk)}</td><td>{e(a.ports)}</td><td>{e(a.services)}</td></tr>"
        for a in assets
    )
    vuln_rows = "\n".join(
        f"<tr><td>{e(v.title)}</td><td>{e(v.cve)}</td><td>{v.cvss:.1f}</td><td>{e(v.severity)}</td><td>{e(v.status)}</td><td>{e(_affected_assets_text(v, asset_map))}</td></tr>"
        for v in vulns
    )
    tracking_rows = "\n".join(
        f"<tr><td>{e(v.title)}</td><td>{e(v.status)}</td><td>{e(v.status_note or '-')}</td><td>{e(v.status_updated_at or '-')}</td><td>{e(v.solution or '-')}</td></tr>"
        for v in vulns
    )
    risk_rows = "\n".join(
        f"<tr><td>{e(item['unit'].name)}</td><td>{item['asset_count']}</td><td>{item['vuln_count']}</td><td>{item['critical_high']}</td><td>{item['remediation']}</td><td>{item['score']}</td><td>{e(_fmt_dt(item['last_sync']))}</td><td>{e(item['last_task_status'] or '-')}</td></tr>"
        for item in _unit_stats(units, assets, vulns, sync_tasks)
    )
    sync_rows = "\n".join(
        f"<tr><td>{e(item['unit'].name)}</td><td>{item['task_count']}</td><td>{item['success_count']}</td><td>{item['failed_count']}</td><td>{item['running_count']}</td><td>{item['success_rate']}%</td><td>{e(item['latest_task'].status if item['latest_task'] else '-')}</td><td>{e((item['latest_task'].error_detail or item['latest_task'].message) if item['latest_task'] else '-')}</td></tr>"
        for item in _sync_quality_rows(units, sync_tasks)
    )
    rendered_content = _render_template_content(report, unit, units, assets, vulns, sync_tasks, template)
    template_section = ""
    if rendered_content:
        body = "".join(f"<p>{e(line)}</p>" for line in rendered_content.splitlines())
        template_section = f"<h2>模板正文</h2>{body}"
    sections = []
    if report_type in {REPORT_TYPE_RISK, REPORT_TYPE_SYNC}:
        sections.append(
            f"<h2>单位风险排行</h2><table><thead><tr><th>单位</th><th>资产数</th><th>漏洞数</th><th>严重/高危</th><th>待整改</th><th>风险评分</th><th>最近同步</th><th>最近任务</th></tr></thead><tbody>{risk_rows}</tbody></table>"
        )
    if report_type == REPORT_TYPE_SYNC:
        sections.append(
            f"<h2>同步质量</h2><table><thead><tr><th>单位</th><th>任务数</th><th>成功</th><th>失败</th><th>执行中</th><th>成功率</th><th>最近状态</th><th>最近结果</th></tr></thead><tbody>{sync_rows}</tbody></table>"
        )
    if report_type in {REPORT_TYPE_ASSET, REPORT_TYPE_VULN, REPORT_TYPE_RISK}:
        sections.append(
            f"<h2>资产清单</h2><table><thead><tr><th>名称</th><th>IP</th><th>单位</th><th>类型</th><th>风险</th><th>端口</th><th>服务</th></tr></thead><tbody>{asset_rows}</tbody></table>"
        )
    if report_type in {REPORT_TYPE_VULN, REPORT_TYPE_RISK}:
        sections.append(
            f"<h2>漏洞清单</h2><table><thead><tr><th>标题</th><th>CVE</th><th>CVSS</th><th>等级</th><th>状态</th><th>影响资产</th></tr></thead><tbody>{vuln_rows}</tbody></table>"
        )
        sections.append(
            f"<h2>处置跟踪</h2><table><thead><tr><th>漏洞</th><th>状态</th><th>处置备注</th><th>状态更新时间</th><th>修复建议</th></tr></thead><tbody>{tracking_rows}</tbody></table>"
        )
    if not sections:
        sections = [
            f"<h2>资产清单</h2><table><thead><tr><th>名称</th><th>IP</th><th>单位</th><th>类型</th><th>风险</th><th>端口</th><th>服务</th></tr></thead><tbody>{asset_rows}</tbody></table>",
            f"<h2>漏洞清单</h2><table><thead><tr><th>标题</th><th>CVE</th><th>CVSS</th><th>等级</th><th>状态</th><th>影响资产</th></tr></thead><tbody>{vuln_rows}</tbody></table>",
        ]
    path.write_text(
        f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <title>{e(report.title)}</title>
  <style>
    body {{ font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif; margin: 32px; color: #1f2329; }}
    .metrics {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px; margin: 20px 0; }}
    .metric {{ border: 1px solid #d9d9d9; padding: 12px; border-radius: 6px; }}
    .metric strong {{ display: block; font-size: 22px; margin-top: 4px; }}
    table {{ width: 100%; border-collapse: collapse; margin: 16px 0 28px; }}
    th, td {{ border: 1px solid #d9d9d9; padding: 8px 10px; text-align: left; }}
    th {{ background: #f5f5f5; }}
  </style>
</head>
<body>
  <h1>{e(report.title)}</h1>
  <p>生成时间：{datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC</p>
  <p>报表范围：{e(unit.name if unit else '全量单位')}</p>
  <p>报表类型：{e(report_type)}</p>
  <p>模板名称：{e(template.name if template else '-')}</p>
  <h2>核心指标</h2>
  <div class="metrics">
    <div class="metric">资产数量<strong>{len(assets)}</strong></div>
    <div class="metric">漏洞数量<strong>{len(vulns)}</strong></div>
    <div class="metric">严重/高危<strong>{sum(1 for v in vulns if v.severity in ['严重', '高危'])}</strong></div>
    <div class="metric">待整改/整改中/待复测<strong>{sum(1 for v in vulns if v.status in ['待整改', '整改中', '待复测'])}</strong></div>
  </div>
  {template_section}
  {''.join(sections)}
</body>
</html>
""",
        encoding="utf-8",
    )


def _affected_assets_text(vuln: Vulnerability, asset_map: dict[str, Asset]) -> str:
    assets = [asset_map[asset_id] for asset_id in (vuln.asset_ids or []) if asset_id in asset_map]
    if not assets:
        return "-"
    return "；".join(f"{asset.name}({asset.ip})" for asset in assets)
